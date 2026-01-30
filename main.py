from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import openpyxl
from num2words import num2words
import os
from difflib import SequenceMatcher

def main():
# генерирую расширенный алфавит бля обращения к ячейкам excel таблицы
    alphabet = generate_alphabet()
# помещаю в переменную объект эксель таблицы
    sheet = load_office_files()
#прописываю путь к папке с документами содержащими темы курсовых и их отметки
    docx_folder = "./path"
    files = os.listdir(docx_folder)
#c помощью функции find_cell_indexes нахожу индекс ячейки(формат одномерный массив из 2 элементов)и добавляю
#в массив marks_indexes
    marks_indexes = []
    marks_indexes.append(find_cell_indexes(sheet, "Фамилия, инициалы (инициал собственного имени) обучающегося"))
    marks_indexes.append(find_cell_indexes(sheet, "Учебная  практика:"))
    marks_indexes.append(find_cell_indexes(sheet, "Преддипломная"))
    marks_indexes.append(find_cell_indexes(sheet, "Количество часов"))

    print(marks_indexes)
    
    hours_row = marks_indexes[3][0]
    names_row = marks_indexes[0][0]
    temp_index = find_cell_indexes(sheet, 'присвоенная квалификация (разряд)')
    for i in range(marks_indexes[0][0] + 1, marks_indexes[3][0]):
        name = sheet[f"{alphabet[marks_indexes[0][1] - 1]}{i}"].value
        qualification = sheet[f"{alphabet[temp_index[1] - 1]}{i}"].value
        doc_temp = Document()
        # Обработка данных
        context1 = find_and_remove(
            return_of_line(sheet, alphabet, i, range(marks_indexes[0][1], (marks_indexes[1][1] - 1)), hours_row,
                           names_row))
        context2 = create_add_table_context(files, name)
        context3 = find_and_remove(
            return_of_line(sheet, alphabet, i, range(marks_indexes[1][1], (marks_indexes[2][1])), hours_row, names_row))

        table1 = create_table(
            doc_temp,
            ['Наименование учебных предметов, модулей, факультативных занятий', 'Количество учебных часов', 'Отметки'],
            context1
        )
        table2 = create_table(
            doc_temp,
            ['Наименование учебных предметов, модулей, по которым выполнялись курсовые работы (курсовые проекты)',
             'Темы курсовых проектов (курсовых работ)', 'Количество учебных часов', 'Отметки'],
            context2)
        table3 = create_table(
            doc_temp,
            ['Наименование практик', 'Количество учебных часов', 'Отметки'],
            context3
        )

        # Работа с основным документом
        doc = Document('G.docx')
        diploma_name=search_diploma_name_in_docx_file(Document('themesDiploma.docx'), name)

        doc = insert_table_in_template(doc, table1, '{{table1}}')
        doc = insert_table_in_template(doc, table2, '{{additional_table}}')
        doc = insert_table_in_template(doc, table3, '{{table2}}')
        doc = replace_placeholder_text(doc, '{{name}}', name)
        doc = replace_placeholder_text(doc, '{{qualification}}', qualification, False)
        doc = replace_placeholder_text(doc,'diploma_name',diploma_name,False)
        doc.save(f'./generated_files/{name} выписка.docx')

def search_diploma_name_in_docx_file(diploma, name):
    temp_table=diploma.tables[0]
        #ищем индекс столбца с именами
    for a, cell in enumerate(temp_table.rows[0].cells):
        if (similar(str(cell.text), "Ф.И.О. учащегося") > 0.8):
            index_of_col_with_names = a
            break     
    #перебираем имена по индексу
    for b in range(len(temp_table.rows) - 1):
        if (similar(str(temp_table.rows[b].cells[index_of_col_with_names].text.replace('\n', ' ')), name) > 0.8):
            return(temp_table.rows[b].cells[index_of_col_with_names+1].text)
    
    #берём таблицу из документа и достаём из неё название темы диплома в соответствии с именем


def generate_alphabet():
    alphabet = [chr(i) for i in range(65, 91)]
    return alphabet + [f"A{char}" for char in alphabet] + [f"B{char}" for char in alphabet]

def load_office_files():
    wb = openpyxl.load_workbook(filename='1.xlsx')
    return wb['Лист1']

def return_of_line(sheet, alphabet, row, column_range, orow, nrow, *args):
    data = []
    for col in column_range:
        temp = [
            sheet[f"{alphabet[col]}{nrow}"].value,
            sheet[f"{alphabet[col]}{orow}"].value,
            sheet[f"{alphabet[col]}{row}"].value
        ]
        data.append(temp)
    return data

def insert_table_in_template(doc, table, placeholder='{{table}}'):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            p = paragraph._p
            parent = p.getparent()
            parent.replace(p, table._tbl)
            return doc
    raise ValueError(f"Плейсхолдер '{placeholder}' не найден")

def create_table(document, headers, rows):
    if not isinstance(document, DocxDocument):
        raise TypeError("Ожидается объект Document")
    # Создание таблицы
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Настройка ширины таблицы и колонок


    if(len(table.columns)==3):
        table.columns[0].width = Inches(4.35)
        table.columns[1].width = Inches(1.25)
        table.columns[2].width = Inches(1.25)
    elif(len(table.columns)==4):
        table.columns[0].width = Inches(3.10)
        table.columns[1].width = Inches(1.25)
        table.columns[2].width = Inches(1.25)
        table.columns[3].width = Inches(1.25)
    # Заголовки
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_format(hdr_cells[i], bold=False)
    # Данные
    for row in rows:
        new_cells = table.add_row().cells
        for i, value in enumerate(row):
            new_cells[i].text = str(value)
            if (len(headers)<=3):            
                set_cell_format(new_cells[i], center= (i == 1))
            else:
                set_cell_format(new_cells[i], center=(i == 1 or i == 2))
    set_table_borders(table)
    return table

def set_table_borders(table):
    """Устанавливает границы для всей таблицы"""
    tbl = table._tbl
    tblBorders = OxmlElement('w:tblBorders')
    borders = {
        'top': {'val': 'single', 'sz': '4', 'color': '000000'},
        'left': {'val': 'single', 'sz': '4', 'color': '000000'},
        'bottom': {'val': 'single', 'sz': '4', 'color': '000000'},
        'right': {'val': 'single', 'sz': '4', 'color': '000000'},
        'insideH': {'val': 'single', 'sz': '4', 'color': '000000'},
        'insideV': {'val': 'single', 'sz': '4', 'color': '000000'}
    }
    for border_name, border_props in borders.items():
        element = OxmlElement(f'w:{border_name}')
        for key, value in border_props.items():
            element.set(qn(f'w:{key}'), value)
        tblBorders.append(element)
    tblPr = tbl.xpath('w:tblPr')[0]
    tblPr.append(tblBorders)

def set_cell_format(cell, bold=False, center=True):
    """Форматирование ячейки"""
    # Шрифт
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            run.font.bold = bold
            # Для кириллицы
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    # Вертикальное выравнивание
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def find_and_remove(context):
    ind = None
    for i in range(0, len(context)):
        for j in range(0, len(context[i])):
            if (context[i][0] == 'Производственная практика:'):
                ind = i
                break
            elif (context[i][j] == None):
                context[i][j] = 'зачтено'
        if (is_number(context[i][2])):
            context[i][2] = str(context[i][2]) + '(' + num2words(int(context[i][2]), lang='ru') + ')'
    if (ind != None):
        del (context[ind])
    return context

def replace_placeholder_text(doc, placeholder, replacement, bold=True):
    """Замена плейсхолдера с форматированием"""
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            paragraph.text = paragraph.text.replace(placeholder, '')
            run = paragraph.add_run(replacement)
            run.font.name = 'Arial'
            run.font.size = Pt(8)
            # Выравнивание
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Для кириллицы
            rPr = run._element.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Arial')
            rFonts.set(qn('w:hAnsi'), 'Arial')
            rPr.append(rFonts)
        # Обработка плейсхолдеров в таблицах

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in cell.text:
                    # Очищаем ячейку и добавляем новый текст
                    cell.text = ''

                    # Добавляем параграф с форматированием
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run(replacement)
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)

                    # Настройки шрифта
                    r = run._element
                    r.rPr.rFonts.set(qn('w:ascii'), 'Arial')
                    r.rPr.rFonts.set(qn('w:hAnsi'), 'Arial')
                    r.rPr.rFonts.set(qn('w:cs'), 'Arial')
                    if (bold):
                        # Выравнивание по центру
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run.font.bold = True
    return doc

def create_add_table_context(files, name):
    data = []
    index_of_col1, index_of_col2, index_of_col3 = 0, 0, 0
    index_of_4 = 0
    for file in files:
        if (file[0] != '~'):
            path_of_file = './path/' + file
        else:
            continue
        doc = Document(path_of_file)
        if (len(doc.tables) != 0):
            our_table = doc.tables[0]
        else:
            continue
        for a, cell in enumerate(our_table.rows[0].cells):
            if (similar(str(cell.text), "-1") > 0.8):
                index_of_col1 = a
                break
        for b in range(len(our_table.rows)):
            if (similar(str(our_table.rows[b].cells[index_of_col1].text.replace('\n', ' ')), name) > 0.8):
                index_of_4 = b
                break
        to1, to3, to2 = our_table.rows[index_of_4].cells[index_of_col1 + 1].text.replace('\n', ' '), \
        our_table.rows[index_of_4].cells[index_of_col1 + 2].text.replace('\n', ' '), 20
        if is_number(to3):
            to3 = str(to3) + '(' + num2words(int(to3), lang='ru') + ')'
        temp = [file.rsplit('.docx', 1)[0], to1, to2, to3]   #      context[i][2] = str(context[i][2]) + '(' + num2words(int(context[i][2]), lang='ru') + ')'
        data.append(temp)
    return data

def find_cell_indexes(sheet, target_value, case_sensitive=True):
    found_cells = []
    # Приведение к строке если необходимо
    if not case_sensitive and isinstance(target_value, str):
        target_value = target_value.lower()
    for row in sheet.iter_rows():
        for cell in row:
            if similar(str(target_value), str(cell.value)) > 0.8:  # similar(str(target_value), str(cell.value)) > 0.8
                found_cells.append(cell.row)
                found_cells.append(cell.column)
    return found_cells

def is_number(s):
    try:
        float(s)
        return True
    except BaseException:
        return False

def similar(a, b):
    return SequenceMatcher(None, a, b).ratio()

if __name__ == "__main__":
    main()

import os
from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import openpyxl


def main():
    # Инициализация основных данных
    alphabet = generate_alphabet()
    sheet = load_office_files()

    # Основной цикл обработки студентов
    for i in range(9, 32):
        student_name = sheet[f"{alphabet[1]}{i}"].value
        context = process_student_data(sheet, alphabet, i)

        # Создание документа
        doc = create_student_document(student_name, context)
        doc.save(f'./generated_files/{student_name} выписка.docx')


def process_student_data(sheet, alphabet, row_idx):
    # Обработка основных данных
    context = {
        'tables': [],
        'additional_data': process_additional_data(sheet, row_idx)
    }

    # Основные таблицы
    for range_def, headers in [
        (range(2, 54), ['Наименование предметов...', 'Часы', 'Отметки']),
        (range(55, 62), ['Наименование практик...', 'Часы', 'Отметки'])
    ]:
        data = process_table_data(sheet, alphabet, row_idx, range_def)
        context['tables'].append(create_table(Document(), headers, data))

    # Таблица из файлов
    context['file_table'] = process_files_table(student_name, sheet, row_idx)

    return context


def process_files_table(student_name, sheet, row_idx):
    # Создание таблицы из файлов
    doc_temp = Document()
    table = doc_temp.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # Заголовки
    headers = ['Учебный предмет', 'Тема', 'Количество часов', 'Отметка']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_format(hdr_cells[i], bold=True)

    # Обработка файлов
    folder_path = './docx_files'
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            file_data = process_single_file(os.path.join(folder_path, filename), student_name)
            if file_data:
                add_file_table_row(table, file_data, sheet, row_idx)

    return table


def process_single_file(filepath, student_name):
    doc = Document(filepath)
    data = {'subject': os.path.splitext(os.path.basename(filepath))[0]}

    for table in doc.tables:
        headers = get_table_headers(table)
        if 'ФИО' in headers and 'Тема' in headers and 'Оценка' in headers:
            for row in table.rows[1:]:  # Пропускаем заголовок
                if student_name in row.cells[headers.index('ФИО')].text:
                    data['topic'] = row.cells[headers.index('Тема')].text
                    data['mark'] = row.cells[headers.index('Оценка')].text
                    return data
    return None


def add_file_table_row(table, file_data, sheet, row_idx):
    # Получение часов из Excel
    subject = file_data['subject']
    hours = get_hours_from_excel(sheet, subject, row_idx)

    # Добавление строки
    row_cells = table.add_row().cells
    row_cells[0].text = subject
    row_cells[1].text = file_data['topic']
    row_cells[2].text = str(hours)
    row_cells[3].text = file_data['mark']

    # Форматирование
    for cell in row_cells:
        set_cell_format(cell)


def get_hours_from_excel(sheet, subject, row_idx):
    # Логика поиска часов по предмету и студенту
    # Реализуйте в соответствии с вашей структурой Excel
    return sheet[f"{find_subject_column(sheet, subject)}{row_idx}"].value


def create_student_document(student_name, context):
    doc = Document('G.docx')

    # Вставка основных таблиц
    for i, table in enumerate(context['tables'], 1):
        doc = insert_table_in_template(doc, table, f'{{{{table{i}}}}')

    # Вставка таблицы из файлов
    doc = insert_table_in_template(doc, context['file_table'], '{{files_table}}')

    # Замена плейсхолдеров
    doc = replace_placeholder_text(doc, '{{name}}', student_name)
    doc = replace_placeholder_text(doc, '{{additional}}', context['additional_data'])

    return doc


# Вспомогательные функции (реализуйте по аналогии с предыдущими)
# - get_table_headers()
# - find_subject_column()
# - set_cell_format()
# - insert_table_in_template()
# - replace_placeholder_text()
# - generate_alphabet()
# - load_office_files()
def return_of_line(sheet, alphabet, row, column_range, *args):
    data = []
    for col in column_range:
        temp = [
            sheet[f"{alphabet[col]}8"].value,
            sheet[f"{alphabet[col]}31"].value,
            sheet[f"{alphabet[col]}{row}"].value
        ]
        data.append(temp)
    return data

if __name__ == "__main__":
    main()